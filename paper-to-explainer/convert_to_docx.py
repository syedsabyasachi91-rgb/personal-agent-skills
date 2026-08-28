#!/usr/bin/env python3
"""Convert markdown or PDF to styled DOCX using python-docx and PyMuPDF."""

from __future__ import annotations

import sys
import os
import argparse
import re

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)

try:
    import fitz  # PyMuPDF
except ImportError:
    print("ERROR: PyMuPDF not installed. Run: pip install pymupdf", file=sys.stderr)
    sys.exit(1)


# --- Style constants ---
BODY_FONT = "Calibri"
BODY_SIZE = Pt(11)
CODE_FONT = "Consolas"
CODE_SIZE = Pt(10)
HEADING_COLORS = {
    1: RGBColor(0x2F, 0x54, 0x96),  # Blue
    2: RGBColor(0x2F, 0x54, 0x96),
    3: RGBColor(0x2F, 0x54, 0x96),
}
HEADING_SIZES = {
    1: Pt(16),
    2: Pt(14),
    3: Pt(12),
}


def create_document():
    """Create a new DOCX document with academic margins."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    return doc


def add_heading_styled(doc, text, level=1):
    """Add a styled heading."""
    size = HEADING_SIZES.get(level, Pt(12))
    color = HEADING_COLORS.get(level, RGBColor(0x2F, 0x54, 0x96))
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.size = size
        run.font.color.rgb = color
        run.font.name = BODY_FONT
    return p


def add_paragraph_styled(doc, text=""):
    """Add a styled body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE
    return p


def add_code_block(doc, code_text):
    """Add an indented code block with monospace font."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Add shading via XML
    shading = p._element.get_or_add_pPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F0F0F0',
        qn('w:val'): 'clear',
    })
    shading.append(shd)
    for line in code_text.split('\n'):
        run = p.add_run(line + '\n')
        run.font.name = CODE_FONT
        run.font.size = CODE_SIZE
    return p


def add_math_block(doc, math_text):
    """Add an indented LaTeX math block with monospace font."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Add shading via XML
    shading = p._element.get_or_add_pPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F0F0F0',
        qn('w:val'): 'clear',
    })
    shading.append(shd)
    for line in math_text.split('\n'):
        run = p.add_run(line + '\n')
        run.font.name = CODE_FONT
        run.font.size = CODE_SIZE
    return p


_INLINE_PATTERN = re.compile(r'(`[^`]+`)|(\*\*[^*]+\*\*)|(\*[^*]+\*)')


def tokenize_inline(text):
    """Split text into (content, style) segments. Style is 'normal'|'code'|'bold'|'italic'."""
    segments = []
    last_end = 0
    for match in _INLINE_PATTERN.finditer(text):
        start, end = match.span()
        if start > last_end:
            segments.append((text[last_end:start], 'normal'))
        token = match.group(0)
        if token.startswith('`'):
            segments.append((token[1:-1], 'code'))
        elif token.startswith('**'):
            segments.append((token[2:-2], 'bold'))
        else:
            segments.append((token[1:-1], 'italic'))
        last_end = end
    if last_end < len(text):
        segments.append((text[last_end:], 'normal'))
    if not segments:
        segments.append((text, 'normal'))
    return segments


def add_inline_formatted(doc, text):
    """Add a paragraph with inline code, bold, and italic formatting."""
    segments = tokenize_inline(text)
    p = doc.add_paragraph()
    for content, style in segments:
        if not content:
            continue
        run = p.add_run(content)
        run.font.name = BODY_FONT
        run.font.size = BODY_SIZE
        if style == 'code':
            run.font.name = CODE_FONT
            run.font.size = Pt(10)
        elif style == 'bold':
            run.bold = True
        elif style == 'italic':
            run.italic = True
    return p


def add_table_from_md(doc, header_line, data_lines):
    """Create a Word table from markdown table lines."""
    headers = [cell.strip() for cell in header_line.strip('| ').split('|')]
    rows = []
    for line in data_lines:
        if not line.strip() or re.match(r'\|[\s\-:|]+\|', line):
            continue
        cells = [cell.strip() for cell in line.strip('| ').split('|')]
        rows.append(cells)
    
    if not headers:
        return
    
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.name = BODY_FONT
        run.font.size = BODY_SIZE
    
    # Data rows
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            if c < len(headers):
                cell = table.rows[r + 1].cells[c]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                run.font.name = BODY_FONT
                run.font.size = BODY_SIZE
    
    return table


def add_bullet_list(doc, items):
    """Add a bullet list."""
    paragraphs = []
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.name = BODY_FONT
        run.font.size = BODY_SIZE
        paragraphs.append(p)
    return paragraphs


def add_blockquote(doc, text):
    """Add an indented blockquote."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.italic = True
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE
    run.font.color.rgb = RGBColor(0x50, 0x50, 0x50)


def convert_markdown_to_docx(md_path: str, output_path: str):
    """Convert a markdown file to styled DOCX."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = create_document()
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_buffer = []
    in_math_block = False
    math_buffer = []
    in_table = False
    table_header = None
    table_data = []
    
    while i < len(lines):
        line = lines[i]
        
        # LaTeX display math blocks ($$...$$)
        if line.strip().startswith('$$'):
            if in_math_block:
                add_math_block(doc, '\n'.join(math_buffer))
                math_buffer = []
                in_math_block = False
            else:
                in_math_block = True
            i += 1
            continue
        
        if in_math_block:
            math_buffer.append(line)
            i += 1
            continue
        
        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue
        
        # Table detection
        if '|' in line and line.strip().startswith('|'):
            # Forward-peek: if already in a table, check if this line starts a new table
            # (new header followed by separator = adjacent table, needs flush first)
            if in_table and not re.match(r'\|[\s\-:|]+\|', line):
                next_is_sep = (i + 1 < len(lines) and re.match(r'\|[\s\-:|]+\|', lines[i + 1]))
                if next_is_sep:
                    # Flush current table before starting new one
                    add_table_from_md(doc, table_header, table_data)
                    in_table = False
                    table_header = None
                    table_data = []
            if not in_table:
                in_table = True
                table_header = line
                table_data = []
            else:
                # Check if separator line
                if re.match(r'\|[\s\-:|]+\|', line):
                    i += 1
                    continue
                table_data.append(line)
            i += 1
            # Check if next line ends table
            if i >= len(lines) or ('|' not in lines[i]):
                add_table_from_md(doc, table_header, table_data)
                in_table = False
                table_header = None
                table_data = []
            continue
        
        # Headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            if level <= 3:
                text = line.strip('# \t')
                add_heading_styled(doc, text, level)
            else:
                text = line.strip('# \t')
                add_heading_styled(doc, text, 3)
            i += 1
            continue
        
        # Blockquote
        if line.strip().startswith('>'):
            quote_text = line.strip().lstrip('> ').strip()
            add_blockquote(doc, quote_text)
            i += 1
            continue
        
        # Bullet list
        if re.match(r'^[\s]*[-*]\s', line):
            items = []
            while i < len(lines) and re.match(r'^[\s]*[-*]\s', lines[i]):
                items.append(re.sub(r'^[\s]*[-*]\s+', '', lines[i]))
                i += 1
            add_bullet_list(doc, items)
            continue
        
        # Numbered list
        if re.match(r'^[\s]*\d+\.\s', line):
            items = []
            while i < len(lines) and re.match(r'^[\s]*\d+\.\s', lines[i]):
                items.append(re.sub(r'^[\s]*\d+\.\s+', '', lines[i]))
                i += 1
            for item in items:
                p = doc.add_paragraph(style='List Number')
                run = p.add_run(item)
                run.font.name = BODY_FONT
                run.font.size = BODY_SIZE
            continue
        
        # Empty line
        if not line.strip():
            i += 1
            continue
        
        # Inline formatting (code, bold, italic) — unified tokenizer
        if _INLINE_PATTERN.search(line):
            add_inline_formatted(doc, line)
        else:
            add_paragraph_styled(doc, line)
        
        i += 1
    
    doc.save(output_path)
    print(f"DOCX saved to: {output_path}")


def convert_pdf_to_docx(pdf_path: str, output_path: str, max_pages: int | None = None):
    """Extract PDF text and create a basic DOCX."""
    doc_pdf = fitz.open(pdf_path)
    total_pages = len(doc_pdf)
    page_end = min(max_pages, total_pages) if max_pages else total_pages
    
    doc = create_document()
    
    for i in range(page_end):
        page = doc_pdf[i]
        text = page.get_text("text")
        
        # Add page heading
        if page_end > 1:
            add_heading_styled(doc, f"Page {i + 1}", level=2)
        
        # Add text as paragraphs
        for line in text.split('\n'):
            line = line.strip()
            if line:
                add_paragraph_styled(doc, line)
    
    doc_pdf.close()
    doc.save(output_path)
    print(f"DOCX saved to: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="Convert markdown or PDF to DOCX")
    parser.add_argument("--input", required=True, help="Input file path (.md or .pdf)")
    parser.add_argument("--output", default=None, help="Output DOCX path (default: same name with .docx)")
    parser.add_argument("--max-pages", type=int, default=None, help="Max pages for PDF input")
    args = parser.parse_args()
    
    input_path = args.input
    if not os.path.exists(input_path):
        print(f"ERROR: File not found: {input_path}", file=sys.stderr)
        sys.exit(1)
    
    if args.output:
        output_path = args.output
    else:
        base = os.path.splitext(input_path)[0]
        output_path = base + ".docx"
    
    ext = os.path.splitext(input_path)[1].lower()
    
    if ext == '.md' or ext == '.markdown':
        convert_markdown_to_docx(input_path, output_path)
    elif ext == '.pdf':
        convert_pdf_to_docx(input_path, output_path, args.max_pages)
    else:
        print(f"ERROR: Unsupported file type: {ext}. Use .md or .pdf", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
